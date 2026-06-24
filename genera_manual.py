"""
Genera Manual_Usuario_NSDP.docx
Ejecutar con: .venv/Scripts/python.exe genera_manual.py
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = "Manual_Usuario_NSDP.docx"
LOGO   = "assets/logo_parroquia.png"

# ── Colores ───────────────────────────────────────────────────────────────────
AZUL_OSCURO  = RGBColor(0x1a, 0x23, 0x7e)
AZUL_MEDIO   = RGBColor(0x1e, 0x40, 0xaf)
GRIS_OSCURO  = RGBColor(0x33, 0x33, 0x33)
GRIS_CLARO   = RGBColor(0x66, 0x66, 0x66)
BLANCO       = RGBColor(0xFF, 0xFF, 0xFF)
AMARILLO_BG  = RGBColor(0xFF, 0xF9, 0xC4)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_margins(doc, top=2.5, bottom=2.5, left=2.5, right=2.5):
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)


def add_page_number(doc):
    """Agrega número de página centrado en el pie de cada sección."""
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.clear()
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run.font.size = Pt(9)
        run.font.color.rgb = GRIS_CLARO


def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = AZUL_OSCURO
    # Línea separadora debajo
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '1a237e')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = GRIS_OSCURO
    return p


def heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = AZUL_MEDIO
    return p


def body(doc, text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(14.5)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    return p


def note(doc, text):
    """Párrafo de nota con fondo amarillo simulado con sombreado."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    # sombreado amarillo
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFF9C4')
    pPr.append(shd)
    run = p.add_run('⚠  ' + text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(0x5a, 0x40, 0x00)
    return p


def table_header_style(cell, text):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = BLANCO
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1a237e')
    tcPr.append(shd)


def table_cell_style(cell, text, bold=False, center=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.clear()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    run.font.name = 'Calibri'


def add_simple_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        table_header_style(t.rows[0].cells[i], h)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table_cell_style(t.rows[ri + 1].cells[ci], val)
    doc.add_paragraph()
    return t


# ── Documento ─────────────────────────────────────────────────────────────────

def build_manual():
    doc = Document()
    set_margins(doc)
    add_page_number(doc)

    # ── PORTADA ───────────────────────────────────────────────────────────────
    # Logo
    if os.path.exists(LOGO):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_before = Pt(60)
        run_logo = p_logo.add_run()
        run_logo.add_picture(LOGO, width=Inches(1.8))
    else:
        p_logo = doc.add_paragraph()
        p_logo.paragraph_format.space_before = Pt(80)

    # Título
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(20)
    run_t = p_title.add_run("Manual de Usuario")
    run_t.bold = True
    run_t.font.size = Pt(28)
    run_t.font.color.rgb = AZUL_OSCURO
    run_t.font.name = 'Calibri'

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s = p_sub.add_run("NSDP — Sistema de Registros Sacramentales  v1.0")
    run_s.font.size = Pt(16)
    run_s.font.color.rgb = AZUL_MEDIO
    run_s.font.name = 'Calibri'

    # Línea decorativa
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_line.paragraph_format.space_before = Pt(10)
    run_l = p_line.add_run("─" * 40)
    run_l.font.color.rgb = AZUL_OSCURO
    run_l.font.size = Pt(12)

    p_par = doc.add_paragraph()
    p_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_par.paragraph_format.space_before = Pt(14)
    run_p = p_par.add_run("Parroquia Nuestra Señora de la Paz")
    run_p.font.size = Pt(14)
    run_p.bold = True
    run_p.font.name = 'Calibri'
    run_p.font.color.rgb = GRIS_OSCURO

    p_loc = doc.add_paragraph()
    p_loc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_loc = p_loc.add_run("Hermosillo, Sonora")
    run_loc.font.size = Pt(12)
    run_loc.font.name = 'Calibri'
    run_loc.font.color.rgb = GRIS_CLARO

    p_fecha = doc.add_paragraph()
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fecha.paragraph_format.space_before = Pt(30)
    run_f = p_fecha.add_run("Junio 2026")
    run_f.font.size = Pt(12)
    run_f.font.name = 'Calibri'
    run_f.font.color.rgb = GRIS_CLARO

    doc.add_page_break()

    # ── TABLA DE CONTENIDO ────────────────────────────────────────────────────
    heading1(doc, "Tabla de Contenido")
    toc_items = [
        ("1.", "Introducción", "3"),
        ("2.", "Pantalla principal", "4"),
        ("3.", "Módulo de Sacramentos", "5"),
        ("   3.1", "Bautismos", "5"),
        ("   3.2", "Primera Comunión", "7"),
        ("   3.3", "Confirmación", "8"),
        ("   3.4", "Matrimonios", "9"),
        ("   3.5", "Catecúmenos", "10"),
        ("4.", "Impresión de formularios", "11"),
        ("5.", "Reportes", "13"),
        ("6.", "Importación desde Excel", "15"),
        ("7.", "Configuración de la Parroquia", "16"),
        ("8.", "Respaldos y Restauración", "17"),
        ("9.", "Preguntas frecuentes", "19"),
    ]
    for num, titulo, pag in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        dots = "." * max(1, 60 - len(num) - len(titulo) - len(pag))
        run = p.add_run(f"{num}  {titulo}  {dots}  {pag}")
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        if not num.startswith(" "):
            run.bold = True

    doc.add_page_break()

    # ── 1. INTRODUCCIÓN ───────────────────────────────────────────────────────
    heading1(doc, "1. Introducción")
    body(doc,
         "NSDP (Sistema de Registros Sacramentales) es una aplicación de escritorio "
         "diseñada para la Parroquia Nuestra Señora de la Paz, Hermosillo, Sonora. "
         "Permite registrar, buscar, reportar e imprimir constancias de los cinco "
         "sacramentos que administra la parroquia.")

    heading2(doc, "¿Para quién es este sistema?")
    body(doc, "Este manual está dirigido al personal de oficina parroquial: secretarias, "
         "administradores y cualquier persona que use NSDP para gestionar los registros "
         "sacramentales. No se requieren conocimientos técnicos de computación avanzados.")

    heading2(doc, "Sacramentos que gestiona")
    add_simple_table(doc,
        ["Sacramento", "Descripción"],
        [
            ["Bautismos",       "Registro de bautizos con datos del bautizado, padres, padrinos y ministro"],
            ["Primera Comunión","Registro de primeras comuniones con nombre, fecha, padres y padrinos"],
            ["Confirmación",    "Registro de confirmaciones con nombre, fecha, arzobispo y padrinos"],
            ["Matrimonios",     "Registro de matrimonios con contrayentes, presbítero y testigos"],
            ["Catecúmenos",     "Registro de catecúmenos con nombre, fecha y padrinos"],
        ]
    )

    heading2(doc, "Requisitos del sistema")
    bullet(doc, "Windows 10 o Windows 11 (64 bits)")
    bullet(doc, "Impresora instalada y configurada (para imprimir constancias y formularios)")
    bullet(doc, "Resolución de pantalla mínima recomendada: 1280 × 720 píxeles")
    note(doc, "No se necesita conexión a Internet. Todos los datos se guardan localmente en la computadora.")

    doc.add_page_break()

    # ── 2. PANTALLA PRINCIPAL ─────────────────────────────────────────────────
    heading1(doc, "2. Pantalla Principal")
    body(doc,
         "Al abrir NSDP aparece la pantalla principal con tres zonas: la cabecera superior, "
         "la barra de pestañas y el área de contenido.")

    heading2(doc, "Cabecera superior")
    body(doc, "La franja superior contiene:")
    bullet(doc, "Título de la aplicación: «NSDP — Registros Sacramentales»")
    bullet(doc, "Botón Reportes — abre la ventana de generación de reportes y exportación")
    bullet(doc, "Botón ⚙ (engrane) — despliega el menú de opciones generales")

    heading2(doc, "Menú ⚙")
    body(doc, "Al hacer clic en el botón ⚙ aparecen tres opciones:")
    bullet(doc, "Datos de la Parroquia — configurar nombre, dirección, logo y demás datos institucionales")
    bullet(doc, "Importar Excel — cargar registros desde archivos Excel (.xlsx)")
    bullet(doc, "Respaldos y Restauración — crear y restaurar copias de seguridad de la base de datos")

    heading2(doc, "Barra de pestañas")
    body(doc,
         "Debajo de la cabecera se encuentra la barra de pestañas. La pestaña «Inicio» muestra "
         "el panel de resumen. Las cinco pestañas siguientes corresponden a cada sacramento: "
         "Bautismos, 1a Comunión, Confirmación, Matrimonios y Catecúmenos.")

    heading2(doc, "Panel de inicio (Dashboard)")
    body(doc,
         "La pantalla de inicio muestra:")
    bullet(doc, "Datos de la parroquia: nombre, dirección, teléfono, párroco actual y redes sociales")
    bullet(doc, "Tarjetas KPI: una por sacramento con el total de registros y los registros del año en curso")
    bullet(doc, "Gráficas de barras con el historial de registros por año para cada sacramento")
    bullet(doc, "Botón «Ver registros» en cada tarjeta para ir directamente a la pestaña correspondiente")

    doc.add_page_break()

    # ── 3. MÓDULO DE SACRAMENTOS ──────────────────────────────────────────────
    heading1(doc, "3. Módulo de Sacramentos")
    body(doc,
         "Cada pestaña de sacramento tiene la misma estructura: una barra de búsqueda, "
         "un listado de registros y botones de acción en la parte inferior.")

    heading2(doc, "Estructura común de cada pestaña")

    heading3(doc, "Barra de búsqueda y filtros")
    bullet(doc, "Campo de texto: filtra el listado en tiempo real mientras escribe el nombre o pareja")
    bullet(doc, "Lista desplegable «Año»: filtra por el año del sacramento (o muestra todos los años)")

    heading3(doc, "Listado de registros")
    body(doc,
         "El listado muestra los registros en páginas de 25. Las columnas se ajustan "
         "automáticamente al contenido. Use la barra de desplazamiento horizontal para "
         "ver columnas adicionales.")
    bullet(doc, "Haga clic en una fila para seleccionar un registro")
    bullet(doc, "Doble clic en una fila para ver todos los campos del registro en una ventana emergente")

    heading3(doc, "Paginación")
    bullet(doc, "Botón «< Anterior»: retrocede una página")
    bullet(doc, "Indicador central: muestra la página actual y el total (p. ej. «Página 2 de 8»)")
    bullet(doc, "Botón «Siguiente >»: avanza una página")
    bullet(doc, "Contador inferior derecho: muestra el total de registros encontrados")

    heading3(doc, "Botones de acción")
    add_simple_table(doc,
        ["Botón", "Color", "Función"],
        [
            ["+ Nuevo",            "Verde",  "Abre el formulario para capturar un nuevo registro"],
            ["Editar",             "Azul",   "Abre el formulario para editar el registro seleccionado (requiere selección)"],
            ["Imprimir constancia","Verde",  "Abre el editor de impresión para el registro seleccionado (requiere selección)"],
        ]
    )

    # ── 3.1 Bautismos ──
    heading2(doc, "3.1 Bautismos")
    body(doc, "Registra los bautizos con información completa del bautizado, sus padres, "
         "padrinos, ministro y datos de registro en libros parroquiales.")

    heading3(doc, "Columnas del listado")
    add_simple_table(doc,
        ["Columna", "Descripción"],
        [
            ["Folio",   "Número secuencial por año, asignado automáticamente"],
            ["Nombre",  "Nombre completo del bautizado"],
            ["Día",     "Día del bautismo"],
            ["Mes",     "Mes del bautismo"],
            ["Año",     "Año del bautismo"],
            ["Libro",   "Número de libro parroquial donde se asentó el acta"],
            ["Página",  "Número de página dentro del libro"],
            ["Acta",    "Número de acta dentro de la página"],
        ]
    )

    heading3(doc, "Campos del formulario de captura")
    add_simple_table(doc,
        ["Campo", "Descripción"],
        [
            ["Nombre",           "Nombre completo del bautizado"],
            ["Día nacimiento",   "Día de nacimiento"],
            ["Mes nacimiento",   "Mes de nacimiento"],
            ["Año nacimiento",   "Año de nacimiento"],
            ["Lugar nacimiento", "Ciudad o municipio de nacimiento"],
            ["Papá",             "Nombre del padre"],
            ["Mamá",             "Nombre de la madre"],
            ["Día bautismo",     "Día del bautismo"],
            ["Mes bautismo",     "Mes del bautismo"],
            ["Año bautismo",     "Año del bautismo"],
            ["Ministro",         "Nombre del ministro que administró el bautismo"],
            ["Padrinos 1",       "Nombres de los primeros padrinos"],
            ["Padrinos 2",       "Nombres de los segundos padrinos (si aplica)"],
            ["Párroco",          "Nombre del párroco"],
            ["Registro No.",     "Número de registro"],
            ["Libro",            "Número de libro"],
            ["Página",           "Número de página"],
            ["Acta",             "Número de acta"],
        ]
    )

    # ── 3.2 Primera Comunión ──
    heading2(doc, "3.2 Primera Comunión")
    body(doc, "Registra las primeras comuniones.")

    heading3(doc, "Columnas del listado")
    body(doc, "Folio, Nombre, Día, Mes, Año.")

    heading3(doc, "Campos del formulario")
    add_simple_table(doc,
        ["Campo", "Descripción"],
        [
            ["Nombre",   "Nombre completo"],
            ["Día",      "Día de la primera comunión"],
            ["Mes",      "Mes"],
            ["Año",      "Año"],
            ["Mamá",     "Nombre de la madre"],
            ["Papá",     "Nombre del padre"],
            ["Padrinos", "Nombres de los padrinos"],
            ["Párroco",  "Nombre del párroco"],
        ]
    )

    # ── 3.3 Confirmación ──
    heading2(doc, "3.3 Confirmación")
    body(doc, "Registra las confirmaciones.")

    heading3(doc, "Columnas del listado")
    body(doc, "Folio, Nombre, Día, Mes, Año, Libro, Página y Acta.")

    heading3(doc, "Campos del formulario")
    add_simple_table(doc,
        ["Campo", "Descripción"],
        [
            ["Número",    "Número interno de confirmación"],
            ["Nombre",    "Nombre completo del confirmado"],
            ["Día",       "Día de la confirmación"],
            ["Mes",       "Mes"],
            ["Año",       "Año"],
            ["Papá",      "Nombre del padre"],
            ["Mamá",      "Nombre de la madre"],
            ["Padrinos",  "Nombres de los padrinos"],
            ["Arzobispo", "Nombre del arzobispo que confirmó"],
            ["Párroco",   "Nombre del párroco"],
            ["Libro",     "Número de libro"],
            ["Página",    "Número de página"],
            ["Acta",      "Número de acta"],
        ]
    )

    # ── 3.4 Matrimonios ──
    heading2(doc, "3.4 Matrimonios")
    body(doc, "Registra los matrimonios.")

    heading3(doc, "Columnas del listado")
    body(doc, "Folio, Pareja, Día, Mes, Año, Libro, Página y Acta.")

    heading3(doc, "Campos del formulario")
    add_simple_table(doc,
        ["Campo", "Descripción"],
        [
            ["Número",      "Número interno de matrimonio"],
            ["Pareja",      "Nombres de los contrayentes"],
            ["Día",         "Día del matrimonio"],
            ["Mes",         "Mes"],
            ["Año",         "Año"],
            ["Presbítero",  "Nombre del presbítero celebrante"],
            ["Testigo 1–4", "Nombres de los testigos (hasta cuatro)"],
            ["Párroco",     "Nombre del párroco"],
            ["Libro",       "Número de libro"],
            ["Página",      "Número de página"],
            ["Acta",        "Número de acta"],
        ]
    )

    # ── 3.5 Catecúmenos ──
    heading2(doc, "3.5 Catecúmenos")
    body(doc, "Registra a los catecúmenos.")

    heading3(doc, "Columnas del listado")
    body(doc, "Folio, Nombre, Día, Mes, Año.")

    heading3(doc, "Campos del formulario")
    add_simple_table(doc,
        ["Campo", "Descripción"],
        [
            ["Nombre",   "Nombre completo"],
            ["Día",      "Día de registro"],
            ["Mes",      "Mes"],
            ["Año",      "Año"],
            ["Padre",    "Nombre del padre"],
            ["Madre",    "Nombre de la madre"],
            ["Padrinos", "Nombres de los padrinos"],
        ]
    )

    heading2(doc, "Crear un nuevo registro")
    numbered(doc, "Vaya a la pestaña del sacramento correspondiente.")
    numbered(doc, "Haga clic en el botón verde «+ Nuevo».")
    numbered(doc, "Complete los campos del formulario.")
    numbered(doc, "Haga clic en «Guardar».")
    note(doc, "El folio se asigna automáticamente al guardar. Es único por año y sacramento.")

    heading2(doc, "Editar un registro existente")
    numbered(doc, "Seleccione el registro en el listado haciendo clic sobre él.")
    numbered(doc, "Haga clic en «Editar».")
    numbered(doc, "Modifique los campos necesarios.")
    numbered(doc, "Haga clic en «Guardar».")

    heading2(doc, "Eliminar un registro")
    numbered(doc, "Seleccione el registro y haga clic en «Editar».")
    numbered(doc, "En el formulario, haga clic en el botón rojo «Eliminar».")
    numbered(doc, "Confirme la eliminación en el cuadro de diálogo.")
    note(doc, "La eliminación no se puede deshacer. Se recomienda crear un respaldo antes de eliminar registros importantes.")

    doc.add_page_break()

    # ── 4. IMPRESIÓN DE FORMULARIOS ───────────────────────────────────────────
    heading1(doc, "4. Impresión de Formularios")
    body(doc,
         "El módulo de impresión permite imprimir los datos de un registro sobre "
         "formularios físicos pre-impresos. Incluye un editor visual donde puede "
         "mover cada campo a la posición exacta requerida por el papel.")

    heading2(doc, "Abrir el editor de impresión")
    numbered(doc, "Vaya a la pestaña del sacramento.")
    numbered(doc, "Seleccione un registro en el listado.")
    numbered(doc, "Haga clic en «Imprimir constancia».")
    body(doc, "Se abre el editor visual de impresión con dos paneles:")
    bullet(doc, "Panel izquierdo: vista previa del formulario con los campos del registro superpuestos")
    bullet(doc, "Panel derecho: tabla de coordenadas X/Y editables para cada campo")

    heading2(doc, "Interfaz del editor visual")

    heading3(doc, "Vista previa")
    bullet(doc, "Muestra la imagen del formulario físico como guía de fondo")
    bullet(doc, "Los campos del registro aparecen en rojo sobre la imagen")
    bullet(doc, "El campo seleccionado se resalta en rojo más intenso con un recuadro")

    heading3(doc, "Controles de zoom")
    bullet(doc, "Botón «−»: reduce el tamaño de la vista previa")
    bullet(doc, "Botón «+»: amplía la vista previa")
    bullet(doc, "Botón «Ajustar»: ajusta la vista al tamaño de la ventana")

    heading3(doc, "Botón «Ocultar guía (solo campos)»")
    body(doc,
         "Alterna entre mostrar la imagen del formulario como fondo y mostrar únicamente "
         "los campos sobre fondo blanco. La vista sin guía muestra exactamente lo que "
         "se imprimirá sobre el papel.")

    heading2(doc, "Mover campos (Drag & Drop)")
    numbered(doc, "Haga clic sobre un campo en la vista previa para seleccionarlo.")
    numbered(doc, "Sin soltar el botón del ratón, arrastre el campo a la nueva posición.")
    numbered(doc, "Suelte el botón para confirmar la posición.")
    body(doc, "Las coordenadas X e Y en el panel derecho se actualizan automáticamente al mover un campo.")

    heading2(doc, "Editar coordenadas manualmente")
    body(doc,
         "También puede escribir directamente los valores de X, Y y tamaño de fuente "
         "en los campos del panel derecho. La vista previa se actualiza en tiempo real.")

    heading2(doc, "Campo «Página» en formularios de Bautismos")
    body(doc,
         "En el formulario físico de bautismo, el campo Página imprime la etiqueta "
         "«Página:» junto con el número de página del libro parroquial. Por ejemplo: "
         "«Página:     255». Este campo es movible igual que los demás.")

    heading2(doc, "Guardar las posiciones")
    body(doc,
         "Al ajustar las posiciones para que coincidan con el formulario físico, "
         "haga clic en «Guardar posiciones». Las posiciones se guardan localmente y "
         "se usarán automáticamente la próxima vez.")
    note(doc, "Las posiciones guardadas son específicas de esta computadora. Si instala NSDP en otro equipo, deberá calibrar las posiciones nuevamente.")

    heading2(doc, "Restablecer posiciones")
    body(doc, "El botón «Restablecer posiciones» regresa todos los campos a las coordenadas predeterminadas del sistema.")

    heading2(doc, "Imprimir")
    numbered(doc, "Una vez que los campos estén en la posición correcta, coloque el formulario físico en la impresora.")
    numbered(doc, "Haga clic en «Imprimir en formulario».")
    numbered(doc, "El sistema enviará el PDF directamente a la impresora predeterminada.")
    note(doc, "Asegúrese de que la impresora esté configurada en orientación vertical (retrato) antes de imprimir.")

    doc.add_page_break()

    # ── 5. REPORTES ───────────────────────────────────────────────────────────
    heading1(doc, "5. Reportes")
    body(doc,
         "El módulo de reportes genera listados consolidados por sacramento, año y párroco. "
         "Acceda desde el botón «Reportes» en la cabecera superior.")

    heading2(doc, "Filtros disponibles")
    add_simple_table(doc,
        ["Filtro", "Descripción"],
        [
            ["Sacramento", "Seleccione el tipo de sacramento (o todos)"],
            ["Año",        "Filtre por año (o todos los años)"],
            ["Mes",        "Filtre por mes (o todos los meses)"],
            ["Párroco",    "Filtre por párroco específico (o todos)"],
        ]
    )

    heading2(doc, "Listado de resultados")
    body(doc,
         "El listado muestra los registros filtrados agrupados por párroco, con subtotales "
         "por grupo y un total general al final.")

    heading2(doc, "Exportar a PDF")
    numbered(doc, "Aplique los filtros deseados.")
    numbered(doc, "Haga clic en «Exportar PDF».")
    numbered(doc, "Seleccione la carpeta y nombre del archivo.")
    body(doc, "El PDF generado incluye:")
    bullet(doc, "Encabezado completo (logo, nombre, dirección y párroco de la parroquia) en cada página")
    bullet(doc, "Agrupación por párroco con conteo por grupo")
    bullet(doc, "Ajuste de línea automático en celdas largas")
    bullet(doc, "Numeración de páginas")
    bullet(doc, "Sección final de Resumen de Totales")

    heading2(doc, "Exportar a Excel")
    numbered(doc, "Aplique los filtros deseados.")
    numbered(doc, "Haga clic en «Exportar Excel».")
    numbered(doc, "Seleccione la carpeta y nombre del archivo.")
    body(doc, "El Excel incluye cortes de grupo por párroco y subtotales.")

    heading2(doc, "Auditoría «Sin Párroco»")
    body(doc,
         "El botón «Sin Párroco → Excel» genera un archivo Excel con cuatro hojas "
         "(Bautismos, Matrimonios, 1a Comunión, Confirmación) que contiene todos los "
         "registros donde el campo Párroco está vacío. Las columnas Folio, Año, Mes, "
         "Día y Párroco se resaltan en amarillo para facilitar la identificación.")
    note(doc, "Use esta función para identificar y completar los registros que no tienen párroco asignado.")

    doc.add_page_break()

    # ── 6. IMPORTACIÓN DESDE EXCEL ────────────────────────────────────────────
    heading1(doc, "6. Importación desde Excel")
    body(doc,
         "NSDP puede importar registros masivamente desde archivos Excel (.xlsx). "
         "Acceda desde el menú ⚙ → «Importar Excel».")

    heading2(doc, "Archivos soportados")
    body(doc,
         "El importador detecta automáticamente las hojas del archivo Excel y mapea "
         "las columnas según el nombre de la hoja. Los archivos deben estar en la "
         "carpeta del programa o en la ubicación configurada.")

    heading2(doc, "Pasos para importar")
    numbered(doc, "Abra el menú ⚙ y seleccione «Importar Excel».")
    numbered(doc, "Verifique que los archivos Excel aparezcan con una palomita verde (✓). "
             "Si aparece una tache roja (✗) significa que el archivo no se encontró.")
    numbered(doc, "Haga clic en «Iniciar importación».")
    numbered(doc, "Espere a que el proceso termine. El progreso se muestra en el área de texto.")
    numbered(doc, "Al finalizar, aparecerá el total de registros importados.")
    note(doc, "La importación no elimina registros existentes. Si un registro ya existe, se omite para evitar duplicados. Se recomienda crear un respaldo antes de importar.")

    doc.add_page_break()

    # ── 7. CONFIGURACIÓN DE LA PARROQUIA ──────────────────────────────────────
    heading1(doc, "7. Configuración de la Parroquia")
    body(doc,
         "Acceda desde el menú ⚙ → «Datos de la Parroquia». "
         "Aquí se configuran los datos institucionales que aparecen en las constancias PDF "
         "y en los reportes.")

    heading2(doc, "Campos disponibles")
    add_simple_table(doc,
        ["Campo", "Aparece en"],
        [
            ["Nombre de la Iglesia",  "Constancias PDF, reportes, pantalla de inicio"],
            ["Ciudad",                "Constancias PDF, pantalla de inicio"],
            ["Colonia",               "Pantalla de inicio"],
            ["Dirección",             "Constancias PDF, pantalla de inicio"],
            ["Código Postal",         "Pantalla de inicio"],
            ["Párroco Actual",        "Constancias PDF, pantalla de inicio"],
            ["Teléfono",              "Pantalla de inicio"],
            ["Horario de Oficina",    "Pantalla de inicio"],
            ["Nombre de la Secretaria","Pantalla de inicio"],
            ["Correo electrónico",    "Pantalla de inicio"],
            ["Facebook",              "Pantalla de inicio"],
            ["Instagram",             "Pantalla de inicio"],
        ]
    )

    heading2(doc, "Logotipo de la parroquia")
    body(doc,
         "El logo a color aparece en el encabezado de las constancias PDF y en la pantalla "
         "de inicio. Tamaño recomendado: 200 × 200 píxeles, formato PNG.")
    numbered(doc, "Haga clic en «Seleccionar logo…».")
    numbered(doc, "Busque y seleccione el archivo de imagen.")
    numbered(doc, "La vista previa mostrará el logo seleccionado.")
    numbered(doc, "Haga clic en «Guardar» para aplicar.")

    heading2(doc, "Logo para reportes PDF")
    body(doc,
         "Versión en negro o escala de grises del logo, usada en el encabezado de los "
         "reportes PDF. Si no se configura, se usa el logo de pantalla como alternativa.")

    heading2(doc, "Foto de la parroquia")
    body(doc,
         "Imagen de la fachada o exterior de la parroquia, que aparece en la pantalla "
         "de inicio junto a los datos institucionales.")

    heading2(doc, "Guardar los cambios")
    body(doc, "Haga clic en «Guardar» para aplicar todos los cambios. "
         "La pantalla de inicio se actualizará automáticamente.")

    doc.add_page_break()

    # ── 8. RESPALDOS Y RESTAURACIÓN ───────────────────────────────────────────
    heading1(doc, "8. Respaldos y Restauración")
    body(doc,
         "Se recomienda crear respaldos periódicos de la base de datos para prevenir "
         "pérdida de información. Acceda desde el menú ⚙ → «Respaldos y Restauración».")

    body(doc,
         "La ventana se divide en dos paneles: el panel izquierdo con la lista de "
         "respaldos y el panel derecho con la bitácora de operaciones.")

    heading2(doc, "Crear un respaldo")
    numbered(doc, "Haga clic en «+ Crear respaldo ahora».")
    numbered(doc, "El sistema crea automáticamente una copia con la fecha y hora actual.")
    numbered(doc, "El respaldo aparece en la lista con su nombre, fecha y tamaño.")
    note(doc, "Se recomienda crear un respaldo antes de realizar importaciones masivas o cambios importantes.")

    heading2(doc, "Restaurar desde la lista")
    numbered(doc, "Seleccione el respaldo deseado en la lista.")
    numbered(doc, "Haga clic en «Restaurar».")
    numbered(doc, "Confirme la operación en el cuadro de diálogo.")
    numbered(doc, "Cierre y vuelva a abrir NSDP para que los cambios surtan efecto.")
    note(doc, "La restauración reemplaza la base de datos actual con el respaldo seleccionado. Los cambios realizados después de ese respaldo se perderán.")

    heading2(doc, "Restaurar desde un archivo externo")
    body(doc,
         "Use esta opción para restaurar un respaldo guardado en otra ubicación o "
         "transferido desde otro equipo:")
    numbered(doc, "Haga clic en «📂 Restaurar desde archivo…».")
    numbered(doc, "Busque y seleccione el archivo .db.")
    numbered(doc, "Confirme y reinicie la aplicación.")

    heading2(doc, "Guardar un respaldo como archivo")
    numbered(doc, "Seleccione el respaldo en la lista.")
    numbered(doc, "Haga clic en «⬇ Guardar como…».")
    numbered(doc, "Seleccione la carpeta de destino y guarde.")
    body(doc, "Esta función es útil para transferir respaldos a unidades externas o enviarlos por correo.")

    heading2(doc, "Eliminar un respaldo")
    numbered(doc, "Seleccione el respaldo en la lista.")
    numbered(doc, "Haga clic en «Eliminar».")
    numbered(doc, "Confirme la eliminación.")
    note(doc, "La eliminación de un respaldo es permanente y no se puede deshacer.")

    heading2(doc, "Bitácora de operaciones")
    body(doc,
         "El panel derecho muestra un historial de todas las operaciones realizadas "
         "(crear, restaurar, eliminar). Cada entrada muestra la fecha y hora, el tipo "
         "de operación, el nombre del archivo y el estado (OK o ERROR).")
    body(doc, "Use el botón «Limpiar» para borrar todos los registros de la bitácora.")

    doc.add_page_break()

    # ── 9. PREGUNTAS FRECUENTES ───────────────────────────────────────────────
    heading1(doc, "9. Preguntas Frecuentes")

    heading2(doc, "¿Cómo busco un registro específico?")
    body(doc,
         "Vaya a la pestaña del sacramento correspondiente. Escriba el nombre (o parte "
         "del nombre) en el campo de búsqueda. El listado se filtra automáticamente "
         "mientras escribe. También puede filtrar por año usando la lista desplegable.")

    heading2(doc, "¿Por qué el folio dice «(asignado al guardar)» al crear un registro?")
    body(doc,
         "El folio es un número secuencial que el sistema asigna automáticamente al "
         "guardar el registro. Se calcula en función del año del sacramento y los folios "
         "ya existentes para ese año.")

    heading2(doc, "¿Los campos Libro, Página y Acta son obligatorios?")
    body(doc,
         "No, son opcionales. Sin embargo, es recomendable capturarlos para facilitar "
         "la búsqueda física del acta en los libros parroquiales y para que aparezcan "
         "en la constancia impresa.")

    heading2(doc, "¿Qué hago si la impresión no queda alineada con el formulario físico?")
    body(doc,
         "Abra el editor de impresión («Imprimir constancia»), active la guía de fondo "
         "y arrastre cada campo hasta que coincida con el espacio correspondiente en el "
         "formulario. Una vez calibrado, haga clic en «Guardar posiciones» para que el "
         "sistema recuerde las coordenadas.")

    heading2(doc, "¿Puedo usar NSDP en más de una computadora?")
    body(doc,
         "Sí. Instale NSDP en cada equipo. Para compartir los datos, exporte un respaldo "
         "desde el equipo principal (⚙ → Respaldos → Guardar como…) y restáurelo en los "
         "demás equipos. Los datos de calibración de posiciones de impresión son "
         "independientes en cada equipo.")

    heading2(doc, "¿Cómo actualizo el nombre del párroco en todos los registros?")
    body(doc,
         "El sistema normaliza automáticamente las variantes tipográficas de los párrocos "
         "históricos. Para registros futuros, ingrese el nombre correcto en el campo "
         "Párroco al capturar cada registro, o actualícelo individualmente usando el "
         "formulario de edición.")

    heading2(doc, "¿Qué hago si la aplicación no abre o muestra un error?")
    bullet(doc, "Verifique que el archivo NSDP.exe esté en la misma carpeta donde se instaló.")
    bullet(doc, "Asegúrese de que la carpeta data\\ exista junto al ejecutable y no esté bloqueada.")
    bullet(doc, "Si el problema persiste, contacte al administrador del sistema o al soporte técnico.")

    heading2(doc, "¿Puedo recuperar un registro eliminado por error?")
    body(doc,
         "Si existe un respaldo previo a la eliminación, puede restaurarlo desde "
         "⚙ → Respaldos y Restauración. Por eso se recomienda crear respaldos frecuentes.")

    # ── Guardar ───────────────────────────────────────────────────────────────
    doc.save(OUTPUT)
    size = os.path.getsize(OUTPUT)
    print(f"Manual generado: {OUTPUT}  ({size:,} bytes)")


if __name__ == "__main__":
    build_manual()
